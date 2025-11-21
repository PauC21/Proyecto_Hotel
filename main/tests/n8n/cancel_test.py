#cancel_test.py
from flask import Flask, jsonify, request, send_from_directory, send_file
import os, time, json
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

app = Flask(__name__)

# ================================
# CONFIGURACIÓN
# ================================
BASE_URL = "http://127.0.0.1:8000"
LOGIN_URL = f"{BASE_URL}/accounts/login/"
RESERVATION_DETAIL_URL = f"{BASE_URL}/reservation/"
GUEST_DETAIL_URL = f"{BASE_URL}/guests/"

USERNAME = "palis"
PASSWORD = "Proyecto2025++"

SCREENSHOT_DIR = os.path.join(os.path.dirname(__file__), "screenshots")
REPORTS_DIR = os.path.join(os.path.dirname(__file__), "reports")

os.makedirs(SCREENSHOT_DIR, exist_ok=True)
os.makedirs(REPORTS_DIR, exist_ok=True)

def safe_text(data):
    if isinstance(data, (dict, list)):
        return json.dumps(data, ensure_ascii=False)
    return str(data)

def take_screenshot(driver, name):
    path = os.path.join(SCREENSHOT_DIR, name)
    driver.save_screenshot(path)
    return path

def generate_word_report(data, filename, images=None):
    from docx import Document
    from docx.shared import Inches

    reports_dir = os.path.join(os.path.dirname(__file__), "reports")
    os.makedirs(reports_dir, exist_ok=True)

    path = os.path.join(reports_dir, filename)
    doc = Document()

    doc.add_heading("Reporte de Prueba Automática", level=1)

    # AGREGAR INFORMACIÓN
    doc.add_paragraph(f"Reserva: {data.get('reservation_id')}")
    doc.add_paragraph(f"Huésped: {data.get('guest_id')}")
    doc.add_paragraph(f"Resultado login: {data.get('login')}")
    doc.add_paragraph(f"Resultado reserva: {data.get('reservation_access')}")
    doc.add_paragraph(f"Cancelación: {data.get('cancel_result')}")
    doc.add_paragraph(f"Validación huésped: {data.get('guest_validation')}")

    # ERRORES
    if data.get("errors"):
        doc.add_heading("Errores", level=2)
        for err in data["errors"]:
            doc.add_paragraph(err)

    doc.add_heading("Capturas", level=2)

    # 🔥 CORREGIDO: procesar solo rutas
    if images:
        for key, img_path in images.items():
            if img_path and os.path.exists(img_path):
                doc.add_paragraph(f"Captura: {key}")
                doc.add_picture(img_path, width=Inches(5.5))
            else:
                doc.add_paragraph(f"[No encontrada: {img_path}]")
    else:
        doc.add_paragraph("Sin capturas")

    doc.save(path)
    return path

#================================
# ENDPOINT: SALUD
# ================================
@app.get("/health")
def health():
    return jsonify({"status": "OK"})


# ================================
# ENDPOINT: EJECUTAR LA PRUEBA
# ================================
@app.post("/run-cancel-test")
def run_cancel_test():
    """
    Ejecuta TODO el Selenium de tu script original,
    y devuelve un JSON para que n8n lo procese.
    """

    # Permite enviar un número de reserva opcional desde n8n
    reservation_id = request.json.get("reservation_id",80)
    guest_id = request.json.get("guest_id", 7)

    RESERVATION_URL = f"{BASE_URL}/reservation/{reservation_id}"
    GUEST_URL = f"{BASE_URL}/guests/{guest_id}"

    result = {
        "reservation_id": reservation_id,
        "guest_id": guest_id,
        "login": None,
        "reservation_access": None,
        "cancel_result": None,
        "guest_validation": None,
        "screenshots": {},
        "errors": []
    }

    # ========== CONFIGURAR SELENIUM ==========
    options = Options()
    #options.add_argument("--headless=new")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--window-size=1920,1080")

    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
    wait = WebDriverWait(driver, 20)

    try:
        # ========= LOGIN =========
        driver.get(LOGIN_URL)

        wait.until(EC.presence_of_element_located((By.NAME, "username")))
        driver.find_element(By.NAME, "username").send_keys(USERNAME)
        driver.find_element(By.NAME, "password").send_keys(PASSWORD)

        driver.find_element(By.XPATH, "//input[@type='submit' and contains(@value,'Continuar')]").click()
        wait.until(EC.url_contains(BASE_URL))

        result["login"] = "OK"

        # ========= DETALLE RESERVA =========
        driver.get(RESERVATION_URL)
        time.sleep(2)

        result["reservation_access"] = "OK"

        # ========= BUSCAR BOTÓN CANCELAR =========
        cancel_button = None
        xpath_list = [
            "//button[contains(., 'Cancelar')]",
            "//button[contains(., '❌')]",
            "//input[contains(@value, 'Cancelar')]",
            "//a[contains(., 'Cancelar')]",
        ]

        for xp in xpath_list:
            try:
                cancel_button = wait.until(EC.presence_of_element_located((By.XPATH, xp)))
                break
            except:
                continue

        if not cancel_button:
            raise Exception("No se encontró botón cancelar")

        driver.execute_script("arguments[0].scrollIntoView();", cancel_button)
        cancel_button.click()

        # Confirmar pop-up si aparece
        try:
            alert = wait.until(EC.alert_is_present())
            alert.accept()
        except:
            pass

        time.sleep(2)

        html_after = driver.page_source.lower()

        if "cancelada" in html_after:
            result["cancel_result"] = "Reserva cancelada correctamente en la página de detalle"
        else:
            result["cancel_result"] = "No se detectó texto de cancelación en reserva"

        # Screenshot después de cancelar
        result["screenshots"]["cancel"] = take_screenshot(driver, "cancel_reservation.png")

        # ========= VALIDAR DETALLE DE HUÉSPED =========
        driver.get(GUEST_URL)
        time.sleep(2)

        guest_html = driver.page_source.lower()

        if str(reservation_id) in guest_html and "cancelada" in guest_html:
            result["guest_validation"] = "Cancelación confirmada en detalle del huésped"
        else:
            result["guest_validation"] = "No aparece la reserva como cancelada en huésped"

        # Screenshot final
        result["screenshots"]["guest"] = take_screenshot(driver, "guest_validation.png")

    except Exception as e:
        error_msg = f"{type(e).__name__}: {e}"
        result["errors"].append(error_msg)

        result["screenshots"]["error"] = take_screenshot(driver, "error.png")

    finally:
        driver.quit()

    # GENERAR REPORTE
    report_name = f"cancel_report_{int(time.time())}.docx"
    report_path = generate_word_report(result, report_name, images=result["screenshots"])

    result["report_file"] = report_name
    result["report_url"] = f"/reports/{report_name}"

    return jsonify(result)


@app.get("/reports/<filename>")
def download_report(filename):
    reports_dir = os.path.join(os.path.dirname(__file__), "reports")
    path = os.path.join(reports_dir, filename)
    if not os.path.exists(path):
        return {"error": "Archivo no encontrado"}, 404
    return send_file(path, as_attachment=True)

# ================================
# INICIO DE SERVIDOR
# ================================
if __name__ == "__main__":
    app.run(port=5003, debug=True)
