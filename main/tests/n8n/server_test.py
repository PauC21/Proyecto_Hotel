# server_test.py
import os, time, json
from flask import Flask, request, jsonify, send_from_directory
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


app = Flask(__name__)

# --- Ajusta esto a tu entorno ---
BASE_URL = "http://127.0.0.1:8000"
LOGIN_URL = f"{BASE_URL}/accounts/login/"
CREATE_URL = f"{BASE_URL}/reservation/create/"
GUEST_DETAIL_URL_TEMPLATE = f"{BASE_URL}/guests/{{guest_id}}/"

USERNAME = "palis"
PASSWORD = "Proyecto2025++"

GUEST_ID = "7"
ROOM_ID = "17"

SCREENSHOT_DIR = os.path.join(os.path.dirname(__file__), "screenshots")
REPORTS_DIR = os.path.join(os.path.dirname(__file__), "reports")
os.makedirs(SCREENSHOT_DIR, exist_ok=True)
os.makedirs(REPORTS_DIR, exist_ok=True)

def take_screenshot(driver, name):
    path = os.path.join(SCREENSHOT_DIR, name)
    driver.save_screenshot(path)
    return path

def set_date_via_js(driver, element, iso_date):
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];
        el.removeAttribute('min');
        el.value = val;
        el.dispatchEvent(new Event('input', { bubbles: true }));
        el.dispatchEvent(new Event('change', { bubbles: true }));
    """, element, iso_date)

def run_reservation_test(payload):
    """
    payload expected:
    {
      "guest_id": 7,
      "room_id": 17,
      "reservas": [{"checkin":"2025-10-22","checkout":"2025-10-24"}, ...]
    }
    """
    guest_id = payload.get("guest_id",7)
    room_id = str(payload.get("room_id", 17))
    reservas = payload.get("reservas", [])

    options = Options()
    # options.add_argument("--headless=new")  # descomenta si quieres sin UI
    options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
    wait = WebDriverWait(driver, 20)

    resumen = {
        "guest_id": guest_id,
        "room_id": room_id,
        "reservas_intentadas": reservas,
        "creadas": [],
        "errores": [],
        "screenshots": []
    }

    try:
        # Login
        driver.get(LOGIN_URL)
        wait.until(EC.presence_of_element_located((By.NAME, "username")))

        driver.find_element(By.NAME, "username").clear() 
        driver.find_element(By.NAME, "username").send_keys(USERNAME)
        driver.find_element(By.NAME, "password").clear()
        driver.find_element(By.NAME, "password").send_keys(PASSWORD)

        driver.find_element(By.XPATH, "//input[@type='submit']").click()

        wait.until(EC.url_changes(LOGIN_URL))
        if LOGIN_URL in driver.current_url:
            raise Exception("Error: no se pudo iniciar sesión, Django devolvió al login.")   

        wait.until(EC.url_contains(BASE_URL))

        # Crear reservas
        for i, r in enumerate(reservas, start=1):
            checkin = r["checkin"]
            checkout = r["checkout"]
            driver.get(CREATE_URL)
            wait.until(EC.presence_of_element_located((By.ID, "id_guest")))

            # seleccionar guest y room
            Select(driver.find_element(By.ID, "id_guest")).select_by_value(guest_id)
            Select(driver.find_element(By.ID, "id_room")).select_by_value(room_id)

            set_date_via_js(driver, driver.find_element(By.ID, "id_check_in_date"), checkin)
            set_date_via_js(driver, driver.find_element(By.ID, "id_check_out_date"), checkout)

            driver.find_element(By.XPATH, "//button[contains(text(),'Guardar reserva')]").click()
            time.sleep(1)

            # captura para evidencia
            sname = f"reserva_{i}_{checkin}_{checkout}.png"
            spath = take_screenshot(driver, sname)
            resumen["screenshots"].append(spath)

            if "exitosamente" in driver.page_source.lower() or "reserva" in driver.current_url:
                resumen["creadas"].append({"index": i, "checkin": checkin, "checkout": checkout})
            else:
                resumen["errores"].append({"index": i, "checkin": checkin, "checkout": checkout, "msg": "No encontrado texto de confirmación"})

        # Verificar detalle del huésped
        guest_detail_url = GUEST_DETAIL_URL_TEMPLATE.format(guest_id=guest_id)
        driver.get(guest_detail_url)
        time.sleep(2)
        sname = f"detalle_huesped_{guest_id}.png"
        resumen["screenshots"].append(take_screenshot(driver, sname))

        # Buscar si las fechas aparecen en el HTML
        page_html = driver.page_source.lower()
        fechas_esperadas = [r["checkin"] for r in reservas]
        # convertir iso a "22 de octubre de 2025" no estrictamente necesario: comprobación simple
        fechas_presentes = [f for f in fechas_esperadas if f in page_html]

        resumen["fechas_presentes"] = fechas_presentes

    except Exception as e:
        resumen["errores"].append({"exception": str(e)})
    finally:
        driver.quit()

    return resumen

def safe_text(data):
    if isinstance(data, dict) or isinstance(data, list):
        return json.dumps(data, ensure_ascii=False)
    return str(data)



def generate_word_report(test_summary, filename, images=None):
    """
    Genera un reporte Word profesional incluyendo texto, tablas e imágenes.

    :param test_summary: dict con los datos del reporte
    :param filename: nombre del archivo a guardar
    :param images: lista opcional de rutas a imágenes que se insertarán en el reporte
    """

    doc = Document()

    # =======================
    #   PORTADA PROFESIONAL
    # =======================
    title = doc.add_heading("Informe de Prueba Automatizada", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph(f"Fecha de generación: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # =======================
    #   RESUMEN EJECUTIVO
    # =======================
    doc.add_heading("Resumen Ejecutivo", level=1)
    doc.add_paragraph(
        "Este documento presenta el resultado de la ejecución automática de pruebas "
        "relacionadas con la creación de reservas en el sistema. Incluye el detalle de las "
        "reservas intentadas, las que se lograron crear y los errores encontrados durante el proceso."
    )
    doc.add_page_break()

    # =======================
    #   RESERVAS SOLICITADAS
    # =======================
    doc.add_heading("Reservas Solicitadas", level=1)

    reservas_intentadas = test_summary.get("reservas_intentadas", [])

    if reservas_intentadas:
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Check-in"
        hdr_cells[1].text = "Check-out"

        for r in reservas_intentadas:
            row = table.add_row().cells
            row[0].text = safe_text(r.get("checkin", ""))
            row[1].text = safe_text(r.get("checkout", ""))
    else:
        doc.add_paragraph("No se solicitaron reservas.")

    doc.add_page_break()

    # =======================
    #   RESERVAS CREADAS
    # =======================
    doc.add_heading("Reservas Creadas", level=1)

    creadas = test_summary.get("creadas", [])
    if creadas:
        for r in creadas:
            doc.add_paragraph(safe_text(r), style="List Number")
    else:
        doc.add_paragraph("No se crearon reservas.")

    doc.add_page_break()

    # =======================
    #   ERRORES
    # =======================
    doc.add_heading("Errores Encontrados", level=1)

    errores = test_summary.get("errores", [])
    if errores:
        for e in errores:
            doc.add_paragraph(safe_text(e), style="List Bullet")
    else:
        doc.add_paragraph("No se registraron errores en la prueba.")

    doc.add_page_break()

    # =======================
    #   AGREGAR IMÁGENES
    # =======================
    doc.add_heading("Evidencias Visuales", level=1)

    if images:
        for img_path in images:
            if os.path.exists(img_path):
                doc.add_paragraph(f"Imagen: {os.path.basename(img_path)}")
                doc.add_picture(img_path, width=Inches(5.5))
                doc.add_paragraph("")  # espacio
            else:
                doc.add_paragraph(f"[No se encontró la imagen: {img_path}]")
    else:
        doc.add_paragraph("No se adjuntaron imágenes en este reporte.")

    # =======================
    #   GUARDAR ARCHIVO
    # =======================
    path = os.path.join("reports", filename)
    os.makedirs("reports", exist_ok=True)
    doc.save(path)

@app.route("/run_test", methods=["POST"])
def run_test():
    payload = request.get_json()
    # ejecutar la prueba
    summary = run_reservation_test(payload)
    # generar reporte
    report_name = f"report_{int(time.time())}.docx"
    report_path = generate_word_report(summary, filename=report_name)
    # construir respuesta
    response = {
        "summary": summary,
        "report_file": report_name,
        "report_url": f"/reports/{report_name}"
    }
    return jsonify(response)

@app.route("/reports/<path:filename>", methods=["GET"])
def download_report(filename):
    return send_from_directory(REPORTS_DIR, filename, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
