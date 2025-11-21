import os, time
from flask import Flask, request, jsonify
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches

# ------------------------------- CONFIGURACIÓN -------------------------------

BASE_URL = "http://127.0.0.1:8000"
LOGIN_URL = f"{BASE_URL}/accounts/login/"
CREATE_URL = f"{BASE_URL}/reservation/create/"

USERNAME = "palis"
PASSWORD = "Proyecto2025++"

GUEST_ID = "6"
ROOM_ID = "16"

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SCREENSHOT_DIR = os.path.join(BASE_DIR, "screenshots")
REPORT_DIR = os.path.join(BASE_DIR, "reports")

os.makedirs(SCREENSHOT_DIR, exist_ok=True)
os.makedirs(REPORT_DIR, exist_ok=True)

app = Flask(__name__)


# ------------------------------- UTILIDADES ----------------------------------

def take_screenshot(driver, name):
    path = os.path.join(SCREENSHOT_DIR, name)
    driver.save_screenshot(path)
    print(f"[CAPTURA] Guardada en: {path}")
    return path


def set_date_via_js(driver, element, iso_date):
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];
        el.removeAttribute('min');
        el.value = val;
        el.dispatchEvent(new Event('input', { bubbles: true }));
        el.dispatchEvent(new Event('change', { bubbles: true }));
    """, element, iso_date)
    print(f"[DEBUG] Fecha establecida en {element.get_attribute('id')}: {iso_date}")


def safe_text(text):
    return str(text).replace("\n", " ").replace("\r", " ")


# --------------------- GENERAR REPORTE WORD PROFESIONAL -----------------------

def generate_word_report(test_summary, filename):
    doc = Document()

    doc.add_heading("Informe de Prueba Automatizada", level=1)
    doc.add_paragraph(f"Fecha de ejecución: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Tipo de prueba: Creación de Reserva")

    doc.add_heading("Parámetros usados:", level=2)
    doc.add_paragraph(f"Check-in: {test_summary['check_in']}")
    doc.add_paragraph(f"Check-out: {test_summary['check_out']}")
    doc.add_paragraph(f"Huésped (ID): {test_summary['guest_id']}")
    doc.add_paragraph(f"Habitación (ID): {test_summary['room_id']}")

    doc.add_heading("Resultado:", level=2)
    doc.add_paragraph(f"Estado: {test_summary['status']}")
    doc.add_paragraph(f"Mensaje: {test_summary.get('message', '')}")

    if test_summary.get("errors"):
        doc.add_heading("Errores:", level=2)
        for e in test_summary["errors"]:
            doc.add_paragraph(f"- {safe_text(e)}")

    doc.add_heading("Evidencias:", level=2)
    for img in test_summary["screenshots"]:
        if os.path.exists(img):
            doc.add_paragraph(os.path.basename(img))
            try:
                doc.add_picture(img, width=Inches(5.5))
            except Exception as e:
                doc.add_paragraph(f"(Error cargando imagen: {e})")

    output_path = os.path.join(REPORT_DIR, filename)
    doc.save(output_path)

    print(f"[REPORTE] Word generado en: {output_path}")
    return output_path


# ------------------------------- FLASK TEST -----------------------------------

@app.route("/run_create_test", methods=["GET", "POST"])
def run_create_test():

    check_in = request.values.get("check_in")
    check_out = request.values.get("check_out")

    if not check_in or not check_out:
        return jsonify({
            "status": "error",
            "message": "Parámetros faltantes: check_in y check_out son requeridos."
        }), 400

    options = Options()
    # Quitar el comentario si quieres MODO INVISIBLE:
    # options.add_argument("--headless=new")
    options.add_argument("--start-maximized")

    driver = webdriver.Chrome(
        service=Service(ChromeDriverManager().install()),
        options=options
    )
    wait = WebDriverWait(driver, 20)

    resultado = {
        "check_in": check_in,
        "check_out": check_out,
        "guest_id": GUEST_ID,
        "room_id": ROOM_ID,
        "status": "",
        "screenshots": [],
        "errors": [],
        "message": ""
    }

    try:
        # ------------------- LOGIN -------------------
        driver.get(LOGIN_URL)
        wait.until(EC.presence_of_element_located((By.NAME, "username")))

        driver.find_element(By.NAME, "username").send_keys(USERNAME)
        driver.find_element(By.NAME, "password").send_keys(PASSWORD)
        driver.find_element(By.XPATH, "//input[@type='submit']").click()

        wait.until(EC.url_contains(BASE_URL))
        # ------------------- CREACIÓN DE RESERVA -------------------
        driver.get(CREATE_URL)
        wait.until(EC.presence_of_element_located((By.ID, "id_guest")))

        Select(driver.find_element(By.ID, "id_guest")).select_by_value(GUEST_ID)
        Select(driver.find_element(By.ID, "id_room")).select_by_value(ROOM_ID)

        checkin_el = driver.find_element(By.ID, "id_check_in_date")
        checkout_el = driver.find_element(By.ID, "id_check_out_date")

        set_date_via_js(driver, checkin_el, check_in)
        set_date_via_js(driver, checkout_el, check_out)

        driver.find_element(By.XPATH, "//button[contains(text(),'Guardar')]").click()
        time.sleep(2)

        page_html = driver.page_source.lower()

        if "exitosamente" in page_html or "reserva" in driver.current_url:
            resultado["status"] = "success"
            resultado["message"] = "Reserva creada exitosamente."
            resultado["screenshots"].append(take_screenshot(driver, "create_success.png"))
        else:
            resultado["status"] = "fail"
            resultado["message"] = "No se encontró mensaje de éxito."
            resultado["screenshots"].append(take_screenshot(driver, "create_fail.png"))

    except Exception as e:
        resultado["status"] = "error"
        resultado["message"] = "Excepción durante la prueba."
        resultado["errors"].append(str(e))
        resultado["screenshots"].append(take_screenshot(driver, "create_error.png"))

    finally:
        driver.quit()

    # ------------------- GENERAR WORD -------------------
    report_name = f"reporte_crear_reserva_{int(time.time())}.docx"
    report_path = generate_word_report(resultado, report_name)

    resultado["report_path"] = report_path

    return jsonify(resultado)


# ------------------------------- EJECUTAR FLASK -------------------------------

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5050)
