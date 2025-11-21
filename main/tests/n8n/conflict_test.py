import os
import time
from flask import Flask, request, jsonify, send_file
from docx import Document
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager

# ==============================
# CONFIGURACIÓN GENERAL
# ==============================

BASE_URL = "http://127.0.0.1:8000"
LOGIN_URL = f"{BASE_URL}/accounts/login/"
CREATE_URL = f"{BASE_URL}/reservation/create/"

USERNAME = "palis"
PASSWORD = "Proyecto2025++"

CHECKIN = "2025-11-22"
CHECKOUT = "2025-11-24"
GUEST_ID = "6"
ROOM_ID = "16"

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SCREENSHOT_DIR = os.path.join(BASE_DIR, "screenshots")
REPORTS_DIR = os.path.join(BASE_DIR, "reports")

os.makedirs(SCREENSHOT_DIR, exist_ok=True)
os.makedirs(REPORTS_DIR, exist_ok=True)

app = Flask(__name__)

# ==============================
# UTILIDADES
# ==============================

def take_screenshot(driver, filename):
    path = os.path.join(SCREENSHOT_DIR, filename)
    driver.save_screenshot(path)
    return path


def set_date_via_js(driver, element, iso_date):
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];
        el.removeAttribute('min');
        el.value = val;
        el.dispatchEvent(new Event('input', { bubbles: true }));
        el.dispatchEvent(new Event('change', { bubbles: true }));
    """, element, iso_date)


# ==============================
# LÓGICA PRINCIPAL DEL TEST
# ==============================

def run_conflict_test():

    result = {
        "guest_id": GUEST_ID,
        "room_id": ROOM_ID,
        "check_in": CHECKIN,
        "check_out": CHECKOUT,
        "screenshots": [],
        "errors": [],
        "status": "",
        "message": "",
        "report_path": ""
    }

    options = Options()
    # 🔥 Comenta esta línea si quieres ver Chrome abrirse
    # options.add_argument("--headless=new")

    options.add_argument("--start-maximized")

    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
    wait = WebDriverWait(driver, 20)

    try:
        # 1. LOGIN
        driver.get(LOGIN_URL)

        wait.until(EC.presence_of_element_located((By.NAME, "username")))
        driver.find_element(By.NAME, "username").send_keys(USERNAME)
        driver.find_element(By.NAME, "password").send_keys(PASSWORD)
        driver.find_element(By.XPATH, "//input[@type='submit' and contains(@value,'Continuar')]").click()

        wait.until(EC.url_contains(BASE_URL))

        # 2. ABRIR FORMULARIO
        driver.get(CREATE_URL)
        wait.until(EC.presence_of_element_located((By.ID, "id_guest")))

        # 3. LLENAR FORMULARIO
        Select(driver.find_element(By.ID, "id_guest")).select_by_value(GUEST_ID)
        Select(driver.find_element(By.ID, "id_room")).select_by_value(ROOM_ID)

        checkin_el = driver.find_element(By.ID, "id_check_in_date")
        checkout_el = driver.find_element(By.ID, "id_check_out_date")

        set_date_via_js(driver, checkin_el, CHECKIN)
        set_date_via_js(driver, checkout_el, CHECKOUT)

        driver.find_element(By.XPATH, "//button[contains(text(),'Guardar reserva')]").click()

        time.sleep(2)
        page_html = driver.page_source.lower()

        expected_error = "esta habitación ya está reservada o no disponible en esas fechas"

        # RESULTADOS
        if expected_error in page_html:
            result["status"] = "success"
            result["message"] = "Conflicto detectado correctamente."
            shot = take_screenshot(driver, "CP-RF-14-2_OK.png")
            result["screenshots"].append(shot)

        elif "exitosamente" in page_html:
            result["status"] = "fail"
            result["message"] = "La reserva fue creada ERRÓNEAMENTE."
            shot = take_screenshot(driver, "CP-RF-14-2_FAIL.png")
            result["screenshots"].append(shot)
        else:
            result["status"] = "warning"
            result["message"] = "No se encontró el mensaje esperado."
            shot = take_screenshot(driver, "CP-RF-14-2_NO_ENCONTRADO.png")
            result["screenshots"].append(shot)

    except Exception as e:
        result["status"] = "error"
        result["message"] = str(e)
        shot = take_screenshot(driver, "CP-RF-14-2_ERROR.png")
        result["screenshots"].append(shot)
        result["errors"].append(str(e))

    finally:
        driver.quit()

    # ==========================
    # GENERAR DOCUMENTO WORD
    # ==========================

    doc = Document()
    doc.add_heading("CP-RF-14-2 - Validar conflicto de reserva", level=1)

    doc.add_paragraph(f"Resultado: {result['status']}")
    doc.add_paragraph(f"Mensaje: {result['message']}")
    doc.add_paragraph(f"Huésped: {GUEST_ID}")
    doc.add_paragraph(f"Habitación: {ROOM_ID}")
    doc.add_paragraph(f"Fechas: {CHECKIN} → {CHECKOUT}")

    for s in result["screenshots"]:
        doc.add_paragraph(s)

    filename = f"reporte_conflicto_reserva_{int(time.time())}.docx"
    filepath = os.path.join(REPORTS_DIR, filename)

    doc.save(filepath)

    result["report_path"] = filepath

    return result


# ==============================
# ENDPOINT PARA EJECUTAR TEST DESDE N8N
# ==============================

@app.post("/run-conflict-test")
def execute_test():
    result = run_conflict_test()
    return jsonify(result)


# ==============================
# ENDPOINT PARA DESCARGAR REPORTES
# ==============================

@app.get("/reports/<filename>")
def download_report(filename):
    path = os.path.join(REPORTS_DIR, filename)
    if not os.path.exists(path):
        return {"error": "Reporte no encontrado"}, 404
    return send_file(path, as_attachment=True)


# ==============================
# MAIN
# ==============================

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
