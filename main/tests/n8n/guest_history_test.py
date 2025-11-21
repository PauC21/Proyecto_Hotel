from flask import Flask, jsonify, request, send_file
import os, time, json
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches

app = Flask(__name__)

# ================================
# CONFIGURACIÓN
# ================================
BASE_URL = "http://127.0.0.1:8000"
LOGIN_URL = f"{BASE_URL}/accounts/login/"

USERNAME = "palis"
PASSWORD = "Proyecto2025++"

SCREENSHOT_DIR = os.path.join(os.path.dirname(__file__), "screenshots")
REPORTS_DIR = os.path.join(os.path.dirname(__file__), "reports")

os.makedirs(SCREENSHOT_DIR, exist_ok=True)
os.makedirs(REPORTS_DIR, exist_ok=True)


def take_screenshot(driver, name):
    path = os.path.join(SCREENSHOT_DIR, name)
    driver.save_screenshot(path)
    return path


def generate_word_report(data, filename, images=None):
    path = os.path.join(REPORTS_DIR, filename)
    doc = Document()

    doc.add_heading("Reporte de Prueba Automática CP-RF-11-3", level=1)

    doc.add_paragraph(f"Huésped ID: {data.get('guest_id')}")
    doc.add_paragraph(f"Resultado login: {data.get('login')}")
    doc.add_paragraph(f"Acceso a detalle huésped: {data.get('guest_page_status')}")
    doc.add_paragraph(f"Validación tabla: {data.get('table_validation')}")
    doc.add_paragraph(f"Validación campos: {data.get('fields_validation')}")

    if data["errors"]:
        doc.add_heading("Errores detectados", level=2)
        for e in data["errors"]:
            doc.add_paragraph(e)

    doc.add_heading("Capturas", level=2)

    if images:
        for key, img in images.items():
            if img and os.path.exists(img):
                doc.add_paragraph(f"Captura: {key}")
                doc.add_picture(img, width=Inches(6))
            else:
                doc.add_paragraph(f"[No encontrada: {img}]")

    doc.save(path)
    return path


@app.get("/health")
def health():
    return jsonify({"status": "OK"})


# ================================
#   NUEVO ENDPOINT PRINCIPAL
# ================================
@app.post("/run_history_test")
def run_history_test():
    """
    Ejecuta la prueba CP-RF-11-3 completa:
    - login
    - ir a detalle de huésped
    - validar tabla de reservas
    - capturas
    - generar reporte
    """

    guest_id = request.json.get("guest_id", 7)
    GUEST_URL = f"{BASE_URL}/guests/{guest_id}/"

    result = {
        "guest_id": guest_id,
        "login": None,
        "guest_page_status": None,
        "table_validation": None,
        "fields_validation": None,
        "screenshots": {},
        "errors": []
    }

    # ========== SELENIUM ==========
    options = Options()
    # options.add_argument("--headless=new")
    options.add_argument("--start-maximized")

    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
    wait = WebDriverWait(driver, 20)

    try:
        # ===== LOGIN =====
        driver.get(LOGIN_URL)
        wait.until(EC.presence_of_element_located((By.NAME, "username")))
        driver.find_element(By.NAME, "username").send_keys(USERNAME)
        driver.find_element(By.NAME, "password").send_keys(PASSWORD)
        driver.find_element(By.XPATH, "//input[@type='submit' and contains(@value,'Continuar')]").click()
        wait.until(EC.url_contains(BASE_URL))
        result["login"] = "OK"

        # ===== DETALLE HUÉSPED =====
        driver.get(GUEST_URL)
        time.sleep(2)
        result["guest_page_status"] = "OK"

        # ===== BUSCAR TABLA =====
        try:
            table = wait.until(EC.presence_of_element_located((By.XPATH, "//table[contains(.,'Reserva') or contains(.,'Habit')]")))
            driver.execute_script("arguments[0].scrollIntoView();", table)
            time.sleep(1)
            result["table_validation"] = "Tabla encontrada"
        except:
            result["table_validation"] = "No se encontró tabla"
            raise Exception("Tabla no encontrada")

        # ===== VALIDAR CAMPOS =====
        html = driver.page_source.lower()
        campos = ["reserva", "habit", "check", "estado"]
        encontrados = [c for c in campos if c in html]

        result["fields_validation"] = f"Campos encontrados: {', '.join(encontrados)}"

        # ===== Screenshot final =====
        result["screenshots"]["final"] = take_screenshot(driver, "consultar_historico.png")

    except Exception as e:
        error = f"{type(e).__name__}: {e}"
        result["errors"].append(error)
        result["screenshots"]["error"] = take_screenshot(driver, "consultar_historico._error.png")

    finally:
        driver.quit()

    # ===== GENERAR REPORTE =====
    report_name = f"guest_history_report_{int(time.time())}.docx"
    generate_word_report(result, report_name, images=result["screenshots"])

    result["report_file"] = report_name
    result["report_url"] = f"/reports/{report_name}"

    return jsonify(result)


@app.get("/reports/<filename>")
def download_report(filename):
    path = os.path.join(REPORTS_DIR, filename)
    if not os.path.exists(path):
        return {"error": "Reporte no encontrado"}, 404
    return send_file(path, as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
